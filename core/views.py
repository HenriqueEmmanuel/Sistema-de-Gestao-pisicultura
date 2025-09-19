import datetime
import json
from django.http import HttpResponseForbidden, JsonResponse, Http404
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from django.contrib.auth.hashers import make_password, check_password
from .models import ParametrosPersonalizados, Tanque, Transacao, Usuario, HistoricoSensor
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import logout
import re
from django.utils.translation import gettext_lazy as _
from django.views.decorators.http import require_POST
from django.db.models import Sum, Q
from django.db.models.functions import TruncMonth
import locale
from django.utils.timezone import now
from django.utils import timezone
from datetime import timedelta

def sair(request):
    logout(request)
    return redirect('index')

def cfg(request):
    return render (request, 'front/cfg.html')

def str_to_bool(value):
    return value.lower() == 'true'

@require_POST
def salvarcfg(request):
    usuario = request.user
    print("POST recebido:", request.POST)

    usuario.nome_fazenda = request.POST.get('nome_fazenda', '')
    usuario.endereco_fazenda = request.POST.get('endereco_fazenda', '')
    usuario.cidade = request.POST.get('cidade', '')
    usuario.estado = request.POST.get('estado', '')

    usuario.notificacoes_push = request.POST.get('notificacoes_push') == 'true'
    usuario.notificacoes_email = request.POST.get('notificacoes_email') == 'true'
    usuario.alertas_agua = request.POST.get('alerta_agua') == 'true'
    usuario.lembretes_alimentacao = request.POST.get('lembrete_alimentacao') == 'true'

    usuario.temperatura = request.POST.get('temperatura', 'C')
    usuario.formato_data = request.POST.get('formato_data', 'DD/MM/AAAA')
    usuario.moeda = request.POST.get('moeda', 'R$')
    usuario.idioma = request.POST.get('idioma', 'pt-BR')

    usuario.backup_automatico = request.POST.get('backup_auto') == 'true'
    usuario.frequencia_backup = request.POST.get('frequencia_backup', 'Semanal')

    usuario.autenticacao_dois_fatores = request.POST.get('dois_fatores') == 'true'

    usuario.tempo_auto_logout = int(request.POST.get('tempo_auto_logout', 30))

    usuario.save()
    return JsonResponse({'mensagem': 'Configurações salvas com sucesso!'})

def senha_forte(senha):
    return (
        len(senha) >= 6 and
        re.search(r"[A-Z]", senha) and
        re.search(r"[a-z]", senha) and
        re.search(r"\d", senha)
    )
#O TERMO DE USO DO SISTEMA TEM "VENDAS" LÁ
def index(request):
    logout(request)
    next_url = request.GET.get('next', 'dashboard')  
    if request.method == "POST":
        acao = request.POST.get("acao")

        # LOGIN
        if acao == "login":
            email_input = request.POST.get("email")
            senha_input = request.POST.get("senha")

            user = authenticate(request, email=email_input, password=senha_input)

            if user is not None:
                login(request, user)
                return redirect("dashboard")
                
            else:
                messages.error(request, "Email ou senha inválidos.")

        # CADASTRO
        elif acao == "cadastro":
            nome = request.POST.get("nome")
            email = request.POST.get("email")
            telefone = request.POST.get("telefone")
            preferencia = request.POST.get("notif")
            senha = request.POST.get("senha")
            repetir_senha = request.POST.get("repetir_senha")
            termos = request.POST.get("termos") == "on"

            if senha != repetir_senha:
                messages.error(request, "As senhas não coincidem.")
                return redirect("index")
            
            if not senha_forte(senha):
                messages.error(request, "A senha deve conter no mínimo 6 caracteres, incluindo letras maiúsculas, minúsculas e números.")
                return redirect("index")

            if not termos:
                messages.error(request, "Você deve aceitar os termos de uso.")
                return redirect("index")

            if Usuario.objects.filter(email=email).exists():
                messages.error(request, "Email já cadastrado.")
                return redirect("index")

            user = Usuario.objects.create_user(
                username=email,
                email=email,
                password=senha,
                first_name=nome,
                telefone=telefone,
                preferencia_notificacao=preferencia,
                aceitou_termos=termos
            )
            messages.success(request, "Cadastro realizado com sucesso! Faça login.")
            return redirect("index")

    return render(request, "front/index.html")


@login_required
def tank(request):
    usuario = request.user  

    tanques = Tanque.objects.filter(usuario=usuario)

    total = tanques.count()
    ativos = tanques.filter(situacao='Ativo').count()
    manutencao = tanques.filter(situacao='Manutenção').count()

    capacidade_total = 0
    for t in tanques:
        if t.comprimento and t.largura and t.altura:
            capacidade_total += (t.comprimento * t.largura * t.altura) / 1000

    context = {
        'tanques': tanques,
        'tem_tanques': total > 0,
        'total': total,
        'ativos': ativos,
        'manutencao': manutencao,
        'capacidade_total': round(capacidade_total, 2),  
    }
    return render(request, 'front/tank.html', context)





def pagina_erro_404(request):
    return render(request, 'front/pagina_erro_404.html')
######
@login_required
def dashboard(request):
    return render(request, 'front/dashboard.html')  

def gerenciamento_sensor(request):
    return render(request, 'front/gerenciamento_sensor.html')

def historico_sensor(request):
    return render(request, 'front/historico_sensor.html')
    #Adicionar

def analise(request):
    return render(request, 'front/analise.html')

def dashboard_content(request):
    return render(request, 'front/dashboard_content.html')
    
def histo_analise(request):
    return render(request, 'front/histo_analise.html')


@login_required(login_url='/index/')
def cadastro_tanque(request):
    usuario = request.user 
    print("Usuário autenticado?", usuario.is_authenticated)

    if request.method == 'POST':
        try:
            nome = request.POST.get('Nome_tanque')
            comprimento = float(request.POST.get('comprimento'))
            largura = float(request.POST.get('largura'))
            altura = float(request.POST.get('altura'))
            tipo = request.POST.get('tipo')
            especie = request.POST.get('especie_cultivada')
            situacao = request.POST.get('situacao')

            capacidade_maxima = request.POST.get('capacidade_maxima')
            capacidade_maxima = int(capacidade_maxima) if capacidade_maxima else None

            temperatura = request.POST.get('temperatura')
            temperatura = float(temperatura) if temperatura else None

            ph = request.POST.get('ph')
            ph = float(ph) if ph else None

            oxigenio = request.POST.get('oxigenio')
            oxigenio = float(oxigenio) if oxigenio else None

            tds = request.POST.get('tds')
            tds = float(tds) if tds else None

            amonia = request.POST.get('amonia')
            amonia = float(amonia) if amonia else None

            nitrito = request.POST.get('nitrito')
            nitrito = float(nitrito) if nitrito else None

            nitrato = request.POST.get('nitrato')
            nitrato = float(nitrato) if nitrato else None

            dureza_geral = request.POST.get('durezageral')
            dureza_geral = float(dureza_geral) if dureza_geral else None

            dureza_carbonatos = request.POST.get('durezacarbonatos')
            dureza_carbonatos = float(dureza_carbonatos) if dureza_carbonatos else None

            salinidade = request.POST.get('salinidade')
            salinidade = float(salinidade) if salinidade else None

            data_instalacao_str = request.POST.get('data_instalacao')
            data_instalacao = (
                datetime.datetime.strptime(data_instalacao_str, '%Y-%m-%d').date()
                if data_instalacao_str else None
            )

            localizacao = request.POST.get('localizacao') or ''
            fonte_agua = request.POST.get('fonte_agua') or ''

            tanque = Tanque.objects.create(
                usuario=usuario,
                nome=nome,
                comprimento=comprimento,
                largura=largura,
                altura=altura,
                capacidade_maxima=capacidade_maxima,
                tipo=tipo,
                especie_cultivada=especie,
                data_instalacao=data_instalacao,
                localizacao=localizacao,
                fonte_agua=fonte_agua,
                situacao=situacao,
                temperatura=temperatura,
                ph=ph,
                oxigenio=oxigenio,
                tds=tds,
                amonia=amonia,
                nitrito=nitrito,
                nitrato=nitrato,
                dureza_geral=dureza_geral,
                dureza_carbonatos=dureza_carbonatos,
                salinidade=salinidade,
            )
            print("Dados recebidos:", request.POST)

            return JsonResponse({'status': 'sucesso'})

        except Exception as e:
            return JsonResponse({'status': 'erro', 'mensagem': str(e)})

    return render(request, 'front/cadastro_tanque.html')





def detalhes_tanque_json(request, tanque_id):
    try:
        tanque = Tanque.objects.get(id=tanque_id)
    except Tanque.DoesNotExist:
        raise Http404("Tanque não encontrado")

    sensores = []

    data = {
        "nome": tanque.nome,
        "tipo": tanque.tipo,
        "especie_cultivada": tanque.especie_cultivada,
        "comprimento": tanque.comprimento,
        "largura": tanque.largura,
        "altura": tanque.altura,
        "capacidade_maxima": tanque.capacidade_maxima,
        "localizacao": tanque.localizacao,
        "fonte_agua": tanque.fonte_agua,
        "situacao": tanque.situacao,
        "temperatura": tanque.temperatura,
        "ph": tanque.ph,
        "oxigenio": tanque.oxigenio,
        "tds": tanque.tds,
        "amonia": tanque.amonia,
        "nitrito": tanque.nitrito,
        "nitrato": tanque.nitrato,
        "dureza_geral": tanque.dureza_geral,
        "dureza_carbonatos": tanque.dureza_carbonatos,
        "salinidade": tanque.salinidade,
        "sensores": sensores,
    }

    return JsonResponse(data)

#CRIAR O SISTEMA de AUTPO DESCOBERTA DE SENSORES ESP32 NA REDE WIFI DA PESSOA




def deletar_tanque(request, tanque_id):
    if request.method == "POST":
        tanque = get_object_or_404(Tanque, id=tanque_id)
        tanque.delete()
        return JsonResponse({'status': 'ok'})
    return JsonResponse({'error': 'Método não permitido'}, status=405)



@login_required
@csrf_exempt
def atualizar_situacao_tanque(request):
    if request.method == 'POST':
        tanque_id = request.POST.get('id')
        nova_situacao = request.POST.get('situacao')

        try:
            tanque = Tanque.objects.get(id=tanque_id, usuario=request.user)
            tanque.situacao = nova_situacao
            tanque.save()

            tanques = Tanque.objects.filter(usuario=request.user)
            ativos = tanques.filter(situacao='Ativo').count()
            manutencao = tanques.filter(situacao='Manutenção').count()

            return JsonResponse({
                'success': True,
                'ativos': ativos,
                'manutencao': manutencao
            })
        except Tanque.DoesNotExist:
            return JsonResponse({'success': False, 'erro': 'Tanque não encontrado.'})
    return JsonResponse({'success': False, 'erro': 'Método inválido'})







def contadores_sensores(request):
    return JsonResponse({
        'total': 0,
        'ativos': 0,
        'esp32': 0,
        'alertas': 0
    })

def escanear_rede(request):
    dispositivos = [
        {'ip': '192.168.0.101'},
        {'ip': '192.168.0.102'},
    ]
    return JsonResponse({'dispositivos': dispositivos})




@login_required
def tanques_do_usuario(request):
    tanques = Tanque.objects.filter(usuario=request.user).values("id", "nome")
    return JsonResponse(list(tanques), safe=False)

@login_required
def dados_sensores(request):
    parametro = request.GET.get("parametro")
    dias = int(request.GET.get("dias", 7))
    tanque_id = request.GET.get("tanque_id")

    if not parametro or not tanque_id:
        return JsonResponse({"error": "Parâmetro ou tanque não informado"}, status=400)

    try:
        data_limite = timezone.now() - timedelta(days=dias)

        dados = HistoricoSensor.objects.filter(
            parametro=parametro,
            tanque__id=tanque_id,
            tanque__usuario=request.user,
            data_hora__gte=data_limite
        ).order_by("data_hora")

        lista = [
            {
                "date": dado.data_hora.strftime("%d/%m/%Y %H:%M"),
                "value": dado.valor
            }
            for dado in dados
        ]

        return JsonResponse({"dados": lista})
    
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return JsonResponse({"error": str(e)}, status=500)


def salvar_parametros_personalizados(request, tanque_id):
    if request.method == "POST":
        try:
            dados = json.loads(request.body.decode('utf-8'))
            Tanque.objects.filter(id=tanque_id).update(**dados)
            return JsonResponse({"sucesso": True})
        except Exception as e:
            return JsonResponse({"sucesso": False, "erro": str(e)})
    return JsonResponse({"sucesso": False, "erro": "Método inválido"})
def get_parametros_personalizados(request, tanque_id):
    try:
        tanque = Tanque.objects.get(id=tanque_id)
        parametros, _ = ParametrosPersonalizados.objects.get_or_create(tanque=tanque)
        parametros = {
            "phminimo": parametros.ph_minimo,
            "phmaximo": parametros.ph_maximo,
            "amoniaminima": parametros.amonia_minima,
            "amoniamaxima": parametros.amonia_maxima,
            "temperaturaminima": parametros.temperatura_minima,
            "temperaturamaxima": parametros.temperatura_maxima,
            "oxigeniodissolvidominimo": parametros.oxigenio_dissolvido_minimo,
            "oxigeniodissolvidomaximo": parametros.oxigenio_dissolvido_maximo,
            "tdsminimo": parametros.tds_minimo,
            "tdsmaximo": parametros.tds_maximo,
            "nitritominimo": parametros.nitrito_minimo,
            "nitritomaximo": parametros.nitrito_maximo,
            "nitratominimo": parametros.nitrato_minimo,
            "nitratomaximo": parametros.nitrato_maximo,
            "dureza1minima": parametros.dureza_geral_minima,
            "dureza1maxima": parametros.dureza_geral_maxima,
            "dureza2minima": parametros.dureza_carbonatos_minima,
            "dureza2maxima": parametros.dureza_carbonatos_maxima,
            "salinidademinima": parametros.salinidade_minima,
            "salinidademaxima": parametros.salinidade_maxima
        }
        return JsonResponse(parametros)
    except Tanque.DoesNotExist:
        return JsonResponse({}, status=404)





def configuracoes_tanque(request, tanque_id):
    tanque = get_object_or_404(Tanque, id=tanque_id)
    
    return render(request, 'front/configuracoes_tanque.html', {
        'tanque': tanque,
        
    })

@csrf_exempt
def salvar_parametros_modal(request, tanque_id):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            tanque = Tanque.objects.get(id=tanque_id)
            parametros, _ = ParametrosPersonalizados.objects.get_or_create(tanque=tanque)

            parametros.ph_minimo = data.get("phminimo") or None
            parametros.ph_maximo = data.get("phmaximo") or None
            parametros.amonia_minima = data.get("amoniaminima") or None
            parametros.amonia_maxima = data.get("amoniamaxima") or None
            parametros.temperatura_minima = data.get("temperaturaminima") or None
            parametros.temperatura_maxima = data.get("temperaturamaxima") or None
            parametros.oxigenio_dissolvido_minimo = data.get("oxigeniodissolvidominimo") or None
            parametros.oxigenio_dissolvido_maximo = data.get("oxigeniodissolvidomaximo") or None
            parametros.tds_minimo = data.get("tdsminimo") or None
            parametros.tds_maximo = data.get("tdsmaximo") or None
            parametros.nitrito_minimo = data.get("nitritominimo") or None
            parametros.nitrito_maximo = data.get("nitritomaximo") or None
            parametros.nitrato_minimo = data.get("nitratominimo") or None
            parametros.nitrato_maximo = data.get("nitratomaximo") or None
            parametros.dureza_geral_minima = data.get("dureza1minima") or None
            parametros.dureza_geral_maxima = data.get("dureza1maxima") or None
            parametros.dureza_carbonatos_minima = data.get("dureza2minima") or None
            parametros.dureza_carbonatos_maxima = data.get("dureza2maxima") or None
            parametros.salinidade_minima = data.get("salinidademinima") or None
            parametros.salinidade_maxima = data.get("salinidademaxima") or None

            parametros.save()
            return JsonResponse({"sucesso": True})
        except Exception as e:
            return JsonResponse({"sucesso": False, "erro": str(e)})
    return JsonResponse({"sucesso": False, "erro": "Método inválido"})

def salvar_configuracoes_tanque(request, tanque_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            tanque = Tanque.objects.get(id=tanque_id)

            tanque.nome = data.get("nome")
            tanque.tipo = data.get("tipo")
            tanque.especie_cultivada = data.get("especie_cultivada")
            tanque.fonte_agua = data.get("fonte_agua")
            tanque.data_instalacao = data.get("data_instalacao") or None
            tanque.localizacao = data.get("localizacao")
            tanque.capacidade_maxima = data.get("capacidade_maxima") or None
            tanque.comprimento = data.get("comprimento") or None
            tanque.largura = data.get("largura") or None
            tanque.altura = data.get("altura") or None
            tanque.situacao = data.get("situacao")
            tanque.temperatura = data.get("temperatura") or None
            tanque.ph = data.get("ph") or None
            tanque.oxigenio = data.get("oxigenio") or None
            tanque.tds = data.get("tds") or None
            tanque.amonia = data.get("amonia") or None
            tanque.nitrito = data.get("nitrito") or None
            tanque.nitrato = data.get("nitrato") or None
            tanque.dureza_geral = data.get("dureza_geral") or None
            tanque.dureza_carbonatos = data.get("dureza_carbonatos") or None
            tanque.salinidade = data.get("salinidade") or None

            tanque.save()

            return JsonResponse({"sucesso": True})
        except Exception as e:
            return JsonResponse({"sucesso": False, "erro": str(e)})

    return JsonResponse({"sucesso": False, "erro": "Requisição inválida"})


def observacoes(request):
    return render(request, 'front/observacoes.html')


def formatar_brl(valor):
    return f"R$ {valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

def transacoes(request):
    transacoes = request.user.transacoes.select_related('tanque').order_by('-data')
    
    soma_receitas = request.user.transacoes.filter(tipo='receita').aggregate(total=Sum('valor'))['total'] or 0
    soma_despesas = request.user.transacoes.filter(tipo='despesa').aggregate(total=Sum('valor'))['total'] or 0

    lucro = soma_receitas - soma_despesas

    margem_lucro = (lucro / soma_receitas * 100) if soma_receitas > 0 else 0

    context = {
        'transacoes': transacoes,
        'soma_receitas_formatado': formatar_brl(soma_receitas),
        'soma_despesas_formatado': formatar_brl(soma_despesas),
        'lucro_formatado': formatar_brl(lucro),
        'margem_lucro_formatado': f"{margem_lucro:.1f}%",
    }
    return render(request, 'front/transacoes.html', context)


def nova_transacao(request):
    if request.method == 'POST':
        try:
            if request.content_type == 'application/json':
                data = json.loads(request.body.decode('utf-8'))       
            else:
                data = request.POST  

            tipo = data.get('tipo')
            if tipo == 'receita':
                categoria = data.get('categoria-receita') or data.get('categoria')
                descricao = data.get('descricao-receita') or ''
                quantidade = float(data.get('quantidade') or 0)
                preco_unitario = float(data.get('preco_unitario') or 0)
                valor = float(data.get('valor') or 0)
                observacoes = data.get('observacoes-receita') or ''
                subcategoria = None  
            else:  
                categoria = data.get('categoria')
                descricao = data.get('descricao') or ''
                quantidade = 0
                preco_unitario = 0
                valor = float(data.get('valor_despesa') or 0)
                observacoes = data.get('observacoes') or ''
                subcategoria = data.get('subcategoria') or None
            
            if not categoria:
                return JsonResponse({'error': 'Selecione uma categoria antes de salvar.'}, status=400)
            if valor <= 0:
                return JsonResponse({'error': 'O valor deve ser maior que zero.'}, status=400)
            print(type(data))



            tanque_id = data.get('tanque')
            tanque = None
            if tanque_id and tanque_id != "geral" and tanque_id != "":
                tanque = Tanque.objects.get(id=tanque_id, usuario=request.user)

            Transacao.objects.create(
                usuario=request.user,
                tipo=tipo,
                categoria=categoria,
                subcategoria=subcategoria,
                descricao=descricao,
                quantidade=quantidade,
                preco_unitario=preco_unitario,
                valor=valor,
                tanque=tanque,
                observacoes=observacoes,
            )
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': True, 'message': 'Transação salva com sucesso!'})

            return JsonResponse({'error': 'Não foi possível salvar a transação. Por favor, tente novamente.'}, status=400)


        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Método não permitido'}, status=405)

from django.http import HttpResponse
from django.utils.timezone import datetime
import csv
import io
from openpyxl import Workbook
from docx import Document
from .models import Transacao

def exportar_transacoes(request):
    formato = request.GET.get('formato', 'pdf')
    periodo = request.GET.get('periodo', 'mes_atual')
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')

    transacoes = Transacao.objects.filter(usuario=request.user)

    hoje = datetime.today()
    if periodo == 'mes_atual':
        transacoes = transacoes.filter(data__month=hoje.month, data__year=hoje.year)
    elif periodo == 'ano_atual':
        transacoes = transacoes.filter(data__year=hoje.year)
    elif periodo == 'personalizado' and data_inicio and data_fim:
        transacoes = transacoes.filter(data__range=[data_inicio, data_fim])

    
    if formato == 'pdf':
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        y = 750
        c.drawString(50, y, "Histórico de Transações")
        y -= 30
        for t in transacoes:
            linha = f"{t.data.strftime('%d/%m/%Y')} | {t.get_tipo_display()} | {t.get_categoria_display()} | R$ {t.valor}"
            c.drawString(50, y, linha)
            y -= 20
            if y < 50:
                c.showPage()
                y = 750
        c.save()
        buffer.seek(0)
        return HttpResponse(buffer, content_type='application/pdf', headers={'Content-Disposition': 'attachment; filename="transacoes.pdf"'})

    elif formato == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.append(['Data', 'Tipo', 'Categoria', 'Descrição', 'Quantidade', 'Preço Unit.', 'Tanque', 'Valor', 'Observações'])
        for t in transacoes:
            ws.append([
                t.data.strftime('%d/%m/%Y'),
                t.get_tipo_display(),
                t.get_categoria_display(),
                t.descricao,
                t.quantidade,
                t.preco_unitario,
                t.tanque.nome if t.tanque else 'Geral',
                t.valor,
                t.observacoes
            ])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=transacoes.xlsx'
        wb.save(response)
        return response

    elif formato == 'docx':
        doc = Document()
        doc.add_heading('Histórico de Transações', 0)
        table = doc.add_table(rows=1, cols=9)
        hdr_cells = table.rows[0].cells
        headers = ['Data','Tipo','Categoria','Descrição','Quantidade','Preço Unit.','Tanque','Valor','Observações']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for t in transacoes:
            row_cells = table.add_row().cells
            row_cells[0].text = t.data.strftime('%d/%m/%Y')
            row_cells[1].text = t.get_tipo_display()
            row_cells[2].text = t.get_categoria_display()
            row_cells[3].text = t.descricao
            row_cells[4].text = str(t.quantidade or '-')
            row_cells[5].text = str(t.preco_unitario or '-')
            row_cells[6].text = t.tanque.nome if t.tanque else 'Geral'
            row_cells[7].text = str(t.valor)
            row_cells[8].text = t.observacoes or ''
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return HttpResponse(f, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': 'attachment; filename=transacoes.docx'})




def format_brl(valor):
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    return locale.currency(valor, grouping=True)

def fluxo_caixa(request):
    transacoes = request.user.transacoes.annotate(mes=TruncMonth('data')).order_by('mes')

    saldo_acumulado = 0
    dados_mensais = []


    meses = (
        transacoes.values('mes')
        .annotate(
            receitas=Sum('valor', filter=Q(tipo='receita')),
            despesas=Sum('valor', filter=Q(tipo='despesa')),
        )
        .order_by('mes')
    )

    for mes_data in meses:
        receitas = mes_data['receitas'] or 0
        despesas = mes_data['despesas'] or 0
        saldo_mensal = receitas - despesas
        saldo_acumulado += saldo_mensal

        dados_mensais.append({
            'mes': mes_data['mes'].strftime('%B de %Y').capitalize(),
            'receitas': format_brl(receitas),
            'despesas': format_brl(despesas),
            'saldo_mensal': format_brl(saldo_mensal),
            'saldo_acumulado': format_brl(saldo_acumulado),
            'class_despesas': 'red' if despesas > 0 else 'green',
            'class_saldo_mensal': 'green' if saldo_mensal >= 0 else 'red',
            'class_saldo_acumulado': 'green' if saldo_acumulado >= 0 else 'red',
            
        })

    return render(request, 'front/fluxo_caixa.html', {
        'dados_mensais': dados_mensais
    })




def format_brl(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def relatorios(request):
    hoje = now().date()
    ano_atual = hoje.year
    mes_atual = hoje.month

    transacoes_ano = Transacao.objects.filter(
        usuario=request.user,
        data__year=ano_atual
    )

    receita_bruta = transacoes_ano.filter(tipo='receita').aggregate(total=Sum('valor'))['total'] or 0

    custos_producao = transacoes_ano.filter(categoria='custo_producao').aggregate(total=Sum('valor'))['total'] or 0

    despesas_operacionais = transacoes_ano.filter(categoria='despesa_operacional').aggregate(total=Sum('valor'))['total'] or 0

    lucro_liquido = receita_bruta - custos_producao - despesas_operacionais

    margem_liquida = (lucro_liquido / receita_bruta * 100) if receita_bruta > 0 else 0

    transacoes_mes = transacoes_ano.filter(data__month=mes_atual)

    custos_variaveis = transacoes_mes.filter(categoria='custo_variavel').aggregate(total=Sum('valor'))['total'] or 0
    custos_fixos = transacoes_mes.filter(categoria='custo_fixo').aggregate(total=Sum('valor'))['total'] or 0
    despesas_operacionais_mes = transacoes_mes.filter(categoria='despesa_operacional').aggregate(total=Sum('valor'))['total'] or 0

    custos_dict = {
        'Custos Variáveis': custos_variaveis,
        'Custos Fixos': custos_fixos,
        'Despesas Operacionais': despesas_operacionais_mes
    }
    maior_custo_nome = max(custos_dict, key=custos_dict.get)
    maior_custo_valor = custos_dict[maior_custo_nome]

    recomendacao = ""
    if maior_custo_nome == 'Custos Variáveis':
        recomendacao = "Focar na otimização dos custos variáveis, especialmente alimentação."
    elif maior_custo_nome == 'Custos Fixos':
        recomendacao = "Avaliar contratos e infraestrutura para reduzir custos fixos."
    elif maior_custo_nome == 'Despesas Operacionais':
        recomendacao = "Revisar processos internos para reduzir despesas operacionais."

    context = {
        # DRE
        'ano_atual': ano_atual,
        'receita_bruta': format_brl(receita_bruta),
        'custos_producao': format_brl(custos_producao),
        'despesas_operacionais': format_brl(despesas_operacionais),
        'lucro_liquido': format_brl(lucro_liquido),
        'lucro_liquido_sinal': f"+ {format_brl(lucro_liquido)}" if lucro_liquido >= 0 else f"- {format_brl(abs(lucro_liquido))}",
        'margem_liquida': f"{margem_liquida:.2f}%",
        'lucro_liquido_class': 'blue' if lucro_liquido >= 0 else 'red',

        # Análise de Custos
        'custos_variaveis': format_brl(custos_variaveis),
        'custos_fixos': format_brl(custos_fixos),
        'despesas_operacionais_mes': format_brl(despesas_operacionais_mes),
        'maior_custo_nome': maior_custo_nome,
        'maior_custo_valor': format_brl(maior_custo_valor),
        'recomendacao': recomendacao,
    }

    return render(request, 'front/relatorios.html', context)


def indicadores(request):
    return render(request, 'front/indicadores.html')